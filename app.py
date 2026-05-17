import streamlit as st
import pandas as pd
import io
import re
from docx import Document
from agent import analiz_surecini_baslat, MusteriStratejisi, MarkaStratejisi, gorsel_analiz_et

st.set_page_config(page_title="AI Ürün Asistanı", page_icon="🛒", layout="wide")

st.title("🛒 Akıllı E-Ticaret Yorum Analisti")
st.markdown("Bu otonom sistem, müşteri yorumlarını analiz eder ve Agentic Framework ile hedef kitleye özel iş zekası raporları sunar.")
st.divider()

def markdown_to_word(doc, md_text):
    for satir in md_text.split('\n'):
        satir = satir.strip()
        if not satir:
            continue
            
        if satir.startswith('### '):
            p = doc.add_heading(level=3)
            temiz_metin = satir[4:]
        elif satir.startswith('## '):
            p = doc.add_heading(level=2)
            temiz_metin = satir[3:]
        elif satir.startswith('# '):
            p = doc.add_heading(level=1)
            temiz_metin = satir[2:]
        elif satir.startswith('* ') or satir.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            temiz_metin = satir[2:]
        else:
            p = doc.add_paragraph()
            temiz_metin = satir

        parcalar = re.split(r'(\*\*.*?\*\*)', temiz_metin)
        for parca in parcalar:
            if parca.startswith('**') and parca.endswith('**'):
                run = p.add_run(parca[2:-2])
                run.bold = True
            else:
                p.add_run(parca)

df = pd.read_csv("veri.csv", encoding="utf-8")

col1, col2 = st.columns([1, 1.5])

with col1:
    st.subheader("📥 Veri Havuzu ve Analitik")
    
    ortalama_puan = df["Yildiz"].mean()
    st.metric(label="Ortalama Memnuniyet", value=f"{ortalama_puan:.1f} / 5.0")
    st.bar_chart(df["Yildiz"].value_counts().sort_index())
    
    with st.expander("Ham Veri Tablosunu Görüntüle"):
        st.dataframe(df, use_container_width=True, hide_index=True)
    
    st.divider()
    
    st.subheader("✍️ Canlı Yorum Ekle")
    with st.form("yeni_yorum_formu", clear_on_submit=True):
        yeni_kullanici = st.text_input("Kullanıcı Adı")
        yeni_yildiz = st.slider("Yıldız Puanı", 1, 5, 5)
        yeni_yorum = st.text_area("Müşteri Yorumu")
        yuklenen_resim = st.file_uploader("📸 Ürün Fotoğrafı Kanıtı (Opsiyonel)", type=["jpg", "jpeg", "png"])
        submit_button = st.form_submit_button("Veri Tabanına Ekle")
        
        if submit_button and yeni_kullanici and yeni_yorum:
            ek_gorsel_notu = ""
            if yuklenen_resim is not None:
                with st.spinner("🤖 Görsel Dedektif fotoğrafı inceliyor..."):
                    gorsel_raporu = gorsel_analiz_et(yuklenen_resim.getvalue())
                    ek_gorsel_notu = f" [Yapay Zeka Görsel Tespiti: {gorsel_raporu}]"
            
            guvenli_yorum = yeni_yorum.replace(",", " ").replace("\n", " ") + ek_gorsel_notu
            guvenli_isim = yeni_kullanici.replace(",", " ")
            
            yeni_satir = pd.DataFrame([{"Kullanici": guvenli_isim, "Yildiz": yeni_yildiz, "Yorum": guvenli_yorum}])
            guncel_df = pd.concat([df, yeni_satir], ignore_index=True)
            guncel_df.to_csv("veri.csv", index=False, encoding="utf-8")
            
            st.success("Yorum ve görsel kanıt başarıyla eklendi! Tablo güncelleniyor...")
            st.rerun()

with col2:
    st.subheader("🤖 Otonom Rapor Sentezi")
    
    if "uretilen_rapor" not in st.session_state:
        st.session_state.uretilen_rapor = None
    
    secilen_hedef = st.radio(
        "Rapor Hedef Kitlesini Seçin:",
        ["Müşteriler (Satın Alma Rehberi)", "Marka Yöneticisi (Ar-Ge ve Kriz Raporu)"],
        horizontal=True
    )
    
    if st.button("🚀 Agentic Analizi Başlat", type="primary", use_container_width=True):
        with st.spinner(f"Ajanlar verileri işliyor..."):
            try:
                yorumlar_metni = "\n- ".join(df["Yorum"].astype(str).tolist())
                
                if secilen_hedef == "Müşteriler (Satın Alma Rehberi)":
                    aktif_strateji = MusteriStratejisi()
                else:
                    aktif_strateji = MarkaStratejisi()
                
                st.session_state.uretilen_rapor = analiz_surecini_baslat(yorumlar_metni, aktif_strateji)
                
            except Exception as e:
                st.error("⚠️ Sistem şu an yoğun veya Google API kota sınırına takıldı. Lütfen 1 dakika bekleyip tekrar deneyin.")
                with st.expander("Geliştirici Hata Detayı (Tıkla)"):
                    st.code(e)

    if st.session_state.uretilen_rapor:
        st.success("✅ Ajanlar görevini başarıyla tamamladı!")
        st.markdown(st.session_state.uretilen_rapor)
        
        st.divider()
        
        doc = Document()
        doc.add_heading("Yapay Zeka Analiz Raporu", 0)
        markdown_to_word(doc, st.session_state.uretilen_rapor)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        dosya_adi = f"Kurumsal_Rapor_{secilen_hedef.split(' ')[0]}.docx"
        
        st.download_button(
            label="📄 Raporu Word Olarak İndir (.docx)",
            data=buffer,
            file_name=dosya_adi,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )